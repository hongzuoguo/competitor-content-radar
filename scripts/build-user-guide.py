#!/usr/bin/env python3
"""Build the public user guide from one canonical JSON source."""

from __future__ import annotations

import argparse
from copy import copy
import json
import re
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "user-guide-content.json"
OUTPUT_DIRECTORY = ROOT / "release"
REQUIRED_SECTIONS = (
    "安装包内置内容",
    "首次配置",
    "主要功能",
    "推荐使用流程",
    "数据和隐私",
    "已知限制",
    "常见问题",
    "使用声明",
)

# compact_reference_guide tokens, plus the named editorial_cover cover override.
TOKENS = {
    "page_width": Inches(8.5),
    "page_height": Inches(11),
    "margin": Inches(1),
    "header_footer": Inches(0.492),
    "body_size": Pt(11),
    "body_after": Pt(6),
    "body_line": 1.25,
    "h1_size": Pt(16),
    "h1_before": Pt(18),
    "h1_after": Pt(10),
    "list_left": Inches(0.375),
    "list_hanging": Inches(0.188),
    "list_after": Pt(4),
    "list_line": 1.25,
}
BLUE = RGBColor(0x2E, 0x74, 0xB5)
NAVY = RGBColor(0x20, 0x37, 0x48)
GOLD = RGBColor(0x7A, 0x5A, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
DOCX_ARCHIVE_TIMESTAMP = (1980, 1, 1, 0, 0, 0)


def set_run_font(run, name: str = "Microsoft YaHei", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_style(style, *, size, color, before, after, line_spacing) -> None:
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = size
    style.font.color.rgb = color
    style.paragraph_format.space_before = before
    style.paragraph_format.space_after = after
    style.paragraph_format.line_spacing = line_spacing


def configure_styles(document: Document) -> None:
    styles = document.styles
    set_paragraph_style(
        styles["Normal"], size=TOKENS["body_size"], color=RGBColor(0, 0, 0),
        before=Pt(0), after=TOKENS["body_after"], line_spacing=TOKENS["body_line"],
    )
    set_paragraph_style(
        styles["Heading 1"], size=TOKENS["h1_size"], color=BLUE,
        before=TOKENS["h1_before"], after=TOKENS["h1_after"], line_spacing=1.0,
    )
    styles["Heading 1"].font.bold = True
    if "Guide Kicker" not in styles:
        styles.add_style("Guide Kicker", WD_STYLE_TYPE.PARAGRAPH)
    set_paragraph_style(
        styles["Guide Kicker"], size=Pt(10.5), color=GOLD,
        before=Pt(0), after=Pt(18), line_spacing=1.0,
    )


def add_numbering_definition(document: Document, number_format: str, force_new: bool = False) -> str:
    """Resolve a real Word numbering definition, creating one only if needed.

    With force_new=True a fresh numId is always created so an ordered list
    restarts at 1 independently of other ordered lists.
    """
    numbering = document.part.numbering_part._element
    abstract_numbers = numbering.findall(qn("w:abstractNum"))
    abstract_by_id = {item.get(qn("w:abstractNumId")): item for item in abstract_numbers}
    if not force_new:
        for number in numbering.findall(qn("w:num")):
            abstract_id = number.find(qn("w:abstractNumId")).get(qn("w:val"))
            abstract = abstract_by_id.get(abstract_id)
            if abstract is None:
                continue
            level = next((item for item in abstract.findall(qn("w:lvl")) if item.get(qn("w:ilvl")) == "0"), None)
            number_format_element = level.find(qn("w:numFmt")) if level is not None else None
            if number_format_element is not None and number_format_element.get(qn("w:val")) == number_format:
                return number.get(qn("w:numId"))

    existing_abstract_ids = [int(value) for value in abstract_by_id if value and value.isdigit()]
    existing_number_ids = []
    for item in numbering.findall(qn("w:num")):
        value = item.get(qn("w:numId"))
        if value and value.isdigit():
            existing_number_ids.append(int(value))
    abstract_id = str(max(existing_abstract_ids, default=-1) + 1)
    number_id = str(max(existing_number_ids, default=0) + 1)

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level_type)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    format_element = OxmlElement("w:numFmt")
    format_element.set(qn("w:val"), number_format)
    level.append(format_element)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if number_format == "bullet" else "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    paragraph_properties = OxmlElement("w:pPr")
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "540")
    indentation.set(qn("w:hanging"), "271")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), number_id)
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), abstract_id)
    number.append(abstract_reference)
    if force_new:
        # Restart numbering at 1 for this independent ordered list.
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        lvl_override.append(start_override)
        number.append(lvl_override)
    numbering.append(number)
    return number_id


def set_num_pr(paragraph, number_id: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), number_id)
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)


def add_text_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = TOKENS["body_after"]
    paragraph.paragraph_format.line_spacing = TOKENS["body_line"]
    set_run_font(paragraph.add_run(text), size=TOKENS["body_size"])


def add_bullet(document: Document, text: str, number_id: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    set_num_pr(paragraph, number_id)
    paragraph.paragraph_format.left_indent = TOKENS["list_left"]
    paragraph.paragraph_format.first_line_indent = -TOKENS["list_hanging"]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = TOKENS["list_after"]
    paragraph.paragraph_format.line_spacing = TOKENS["list_line"]
    set_run_font(paragraph.add_run(text), size=TOKENS["body_size"])


def add_ordered_item(document: Document, text: str, number_id: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    set_num_pr(paragraph, number_id)
    paragraph.paragraph_format.left_indent = TOKENS["list_left"]
    paragraph.paragraph_format.first_line_indent = -TOKENS["list_hanging"]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = TOKENS["list_after"]
    paragraph.paragraph_format.line_spacing = TOKENS["list_line"]
    set_run_font(paragraph.add_run(text), size=TOKENS["body_size"])


def add_footer(section, presentation: dict) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(presentation["docx_footer_prefix"])
    set_run_font(run, size=Pt(9), color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    run = paragraph.add_run(presentation["docx_footer_suffix"])
    set_run_font(run, size=Pt(9), color=MUTED)


def add_header(section, presentation: dict) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(presentation["docx_header"])
    set_run_font(run, size=Pt(9), color=MUTED)


def add_cover(document: Document, guide: dict) -> None:
    presentation = guide["presentation"]
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(132)
    kicker = document.add_paragraph(style="Guide Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(kicker.add_run(presentation["cover_kicker"]), size=Pt(10.5), color=GOLD, bold=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run(guide["title"]), size=Pt(30), color=NAVY, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_run_font(subtitle.add_run(guide["subtitle"]), size=Pt(15), color=RGBColor(0x2B, 0x51, 0x63))
    statement = document.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.space_after = Pt(88)
    set_run_font(statement.add_run(presentation["cover_notice"]), size=Pt(10.5), color=GOLD)
    audience = document.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(audience.add_run(presentation["cover_audience"]), size=Pt(9.5), color=MUTED, italic=True)
    document.add_page_break()


def validate_guide(guide: dict) -> None:
    if not isinstance(guide, dict):
        raise ValueError("Guide must be a JSON object.")
    expected_top_level_keys = {"title", "subtitle", "output_stem", "presentation", "sections"}
    if set(guide) != expected_top_level_keys:
        raise ValueError("Guide has missing or unexpected top-level fields.")
    for key in ("title", "subtitle", "output_stem"):
        if not isinstance(guide[key], str) or not guide[key].strip():
            raise ValueError(f"Guide field must be a non-empty string: {key}")

    output_stem = guide["output_stem"]
    if (
        output_stem in {".", ".."}
        or "/" in output_stem
        or "\\" in output_stem
        or ".." in Path(output_stem).parts
        or Path(output_stem).is_absolute()
        or re.match(r"^[A-Za-z]:", output_stem)
    ):
        raise ValueError("Guide output stem must be a simple file name.")

    presentation = guide["presentation"]
    required_presentation_keys = {
        "cover_kicker", "cover_notice", "cover_audience", "docx_header",
        "docx_footer_prefix", "docx_footer_suffix",
    }
    if not isinstance(presentation, dict) or set(presentation) != required_presentation_keys:
        raise ValueError("Guide presentation has missing or unexpected fields.")
    for key, value in presentation.items():
        if not isinstance(value, str) or not value.strip():
            raise ValueError(f"Presentation field must be a non-empty string: {key}")

    sections = guide["sections"]
    if not isinstance(sections, list) or not sections:
        raise ValueError("Guide must provide a non-empty sections list.")
    titles = []
    section_content_keys = {"paragraphs", "items", "ordered_items"}
    for section in sections:
        if not isinstance(section, dict) or set(section) - ({"title"} | section_content_keys):
            raise ValueError("Guide section has unexpected fields.")
        title = section.get("title")
        if not isinstance(title, str) or not title.strip():
            raise ValueError("Each section title must be a non-empty string.")
        titles.append(title)
        present_content_keys = section_content_keys & set(section)
        if not present_content_keys:
            raise ValueError(f"Section has no content: {title}")
        for key in present_content_keys:
            values = section[key]
            if not isinstance(values, list) or not values or any(not isinstance(value, str) or not value.strip() for value in values):
                raise ValueError(f"Section field must be a non-empty string list: {title}.{key}")
    missing = [title for title in REQUIRED_SECTIONS if title not in titles]
    if missing:
        raise ValueError(f"Guide is missing required sections: {', '.join(missing)}")


def output_paths(guide: dict, output_directory: Path = OUTPUT_DIRECTORY) -> tuple[Path, Path]:
    output_root = output_directory.resolve()
    markdown_path = (output_root / f"{guide['output_stem']}.md").resolve()
    docx_path = (output_root / f"{guide['output_stem']}.docx").resolve()
    for path in (markdown_path, docx_path):
        try:
            path.relative_to(output_root)
        except ValueError as error:
            raise ValueError("Guide output path must remain inside release.") from error
    return markdown_path, docx_path


def markdown(guide: dict) -> str:
    lines = [f"# {guide['title']}", "", guide["subtitle"], ""]
    for section in guide["sections"]:
        lines.extend([f"## {section['title']}", ""])
        for paragraph in section.get("paragraphs", []):
            lines.extend([paragraph, ""])
        for item in section.get("items", []):
            lines.append(f"- {item}")
        for item in section.get("ordered_items", []):
            lines.append(f"1. {item}")
        if section.get("items") or section.get("ordered_items"):
            lines.append("")
    return "\n".join(lines)


def docx_document(guide: dict, path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = TOKENS["page_width"]
    section.page_height = TOKENS["page_height"]
    section.top_margin = TOKENS["margin"]
    section.bottom_margin = TOKENS["margin"]
    section.left_margin = TOKENS["margin"]
    section.right_margin = TOKENS["margin"]
    section.header_distance = TOKENS["header_footer"]
    section.footer_distance = TOKENS["header_footer"]
    configure_styles(document)
    add_header(section, guide["presentation"])
    add_footer(section, guide["presentation"])
    add_cover(document, guide)
    bullet_number_id = add_numbering_definition(document, "bullet")
    for section_data in guide["sections"]:
        heading = document.add_paragraph(style="Heading 1")
        set_run_font(heading.add_run(section_data["title"]), size=TOKENS["h1_size"], color=BLUE, bold=True)
        for paragraph in section_data.get("paragraphs", []):
            add_text_paragraph(document, paragraph)
        for item in section_data.get("items", []):
            add_bullet(document, item, bullet_number_id)
        if section_data.get("ordered_items"):
            decimal_number_id = add_numbering_definition(document, "decimal", force_new=True)
            for item in section_data["ordered_items"]:
                add_ordered_item(document, item, decimal_number_id)
    document.save(path)
    normalize_docx_archive_timestamps(path)


def normalize_docx_archive_timestamps(path: Path) -> None:
    temporary_path = path.with_name(f".{path.name}.tmp")
    try:
        with ZipFile(path, "r") as source, ZipFile(temporary_path, "w") as destination:
            for entry in source.infolist():
                normalized_entry = copy(entry)
                normalized_entry.date_time = DOCX_ARCHIVE_TIMESTAMP
                destination.writestr(normalized_entry, source.read(entry.filename))
        temporary_path.replace(path)
    finally:
        temporary_path.unlink(missing_ok=True)


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the HitMuse user guide")
    parser.add_argument(
        "--output-directory",
        type=Path,
        default=OUTPUT_DIRECTORY,
        help="Directory for the generated Markdown and DOCX files",
    )
    return parser.parse_args()


def main(output_directory: Path = OUTPUT_DIRECTORY) -> int:
    try:
        guide = json.loads(SOURCE.read_text(encoding="utf-8"))
        validate_guide(guide)
        output_directory.mkdir(parents=True, exist_ok=True)
        markdown_path, docx_path = output_paths(guide, output_directory)
        markdown_path.write_text(markdown(guide), encoding="utf-8", newline="\n")
        docx_document(guide, docx_path)
    except (OSError, ValueError, json.JSONDecodeError) as error:
        print(f"Unable to build user guide from canonical source {SOURCE}: {error}", file=sys.stderr)
        return 1
    for path in (markdown_path, docx_path):
        print(path)
    return 0


if __name__ == "__main__":
    arguments = parse_arguments()
    raise SystemExit(main(arguments.output_directory))
